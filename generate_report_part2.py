from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    
    # Set default font and line spacing
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    return doc

def add_certificate_page(doc):
    print("✅ Adding certificate page...")
    
    # Add certificate heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CERTIFICATE')
    run.bold = True
    run.font.size = Pt(16)
    
    # Add space after heading
    doc.add_paragraph('\n\n')
    
    # Add certificate content
    content = [
        "This is to certify that the project titled",
        "\"SMART MARINE PROJECT: AI-POWERED PLASTIC WASTE DETECTION",
        "AND AUTONOMOUS COLLECTION SYSTEM\"",
        "submitted by",
        "SUDEEKSHA (20000XXXXX)",
        "[Student Name 2] (20000XXXXX)",
        "[Student Name 3] (20000XXXXX)",
        "[Student Name 4] (20000XXXXX)",
        "in partial fulfillment of the requirements for the award of the degree of",
        "BACHELOR OF TECHNOLOGY",
        "in",
        "COMPUTER SCIENCE AND ENGINEERING",
        "is a bonafide record of work carried out by them under my supervision."
    ]
    
    for line in content:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        if line in ["BACHELOR OF TECHNOLOGY", "COMPUTER SCIENCE AND ENGINEERING"]:
            run.bold = True
    
    # Add space for signatures
    for _ in range(10):
        doc.add_paragraph()
    
    # Add signature line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("___________________________")
    
    # Add page break
    doc.add_page_break()
    
    return doc

def add_declaration_page(doc):
    print("✅ Adding declaration page...")
    
    # Add declaration heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DECLARATION')
    run.bold = True
    run.font.size = Pt(16)
    
    # Add space after heading
    doc.add_paragraph('\n\n')
    
    # Add declaration content
    declaration = [
        "We hereby declare that the work which is being presented in the project titled",
        "\"SMART MARINE PROJECT: AI-POWERED PLASTIC WASTE DETECTION",
        "AND AUTONOMOUS COLLECTION SYSTEM\"",
        "is an authentic record of our own work carried out under the supervision of",
        "[Guide Name], [Designation], Department of CSE, ANURAG UNIVERSITY",
        "during the academic year 2023-24.",
        "",
        "The results embodied in this project have not been submitted to any other",
        "University or Institute for the award of any degree or diploma."
    ]
    
    for line in declaration:
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = 1.5
    
    # Add space for signatures
    for _ in range(5):
        doc.add_paragraph()
    
    # Add date and place
    p = doc.add_paragraph("Date: ")
    p.add_run("\t" * 6 + "Place: ")
    
    # Add student signatures
    students = [
        "1. SUDEEKSHA (20000XXXXX)",
        "2. [Student Name 2] (20000XXXXX)",
        "3. [Student Name 3] (20000XXXXX)",
        "4. [Student Name 4] (20000XXXXX)"
    ]
    
    for student in students:
        doc.add_paragraph()
        p = doc.add_paragraph(student)
        p.add_run("\t" * 6 + "___________________________")
    
    # Add guide signature
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("\t" * 6 + "___________________________")
    p = doc.add_paragraph()
    p.add_run("\t" * 7 + "[Guide Name]")
    p = doc.add_paragraph()
    p.add_run("\t" * 7 + "[Designation]")
    p = doc.add_paragraph()
    p.add_run("\t" * 6 + "Department of CSE")
    
    # Add page break
    doc.add_page_break()
    
    return doc

def add_acknowledgment_page(doc):
    print("✅ Adding acknowledgment page...")
    
    # Add acknowledgment heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ACKNOWLEDGMENT')
    run.bold = True
    run.font.size = Pt(16)
    
    # Add space after heading
    doc.add_paragraph('\n\n')
    
    # Add acknowledgment content
    acknowledgment = [
        "We express our sincere gratitude to our guide [Guide Name], [Designation], ",
        "Department of Computer Science and Engineering, ANURAG UNIVERSITY, for their ",
        "valuable guidance, constant encouragement, and constructive suggestions ",
        "throughout the course of this project work.",
        "",
        "We are also thankful to Dr. [HOD Name], Head of the Department of Computer ",
        "Science and Engineering, and all the faculty members of the department for ",
        "their support and encouragement.",
        "",
        "Our sincere thanks to the Management and Principal of ANURAG UNIVERSITY for ",
        "providing the necessary facilities and support for carrying out this project work.",
        "",
        "Last but not least, we would like to express our heartfelt gratitude to our ",
        "parents and friends for their constant support and motivation throughout the ",
        "completion of this project."
    ]
    
    for line in acknowledgment:
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = 1.5
    
    # Add page break
    doc.add_page_break()
    
    return doc

def main():
    print("📄 Starting report generation (Part 2)...")
    doc = create_document()
    doc = add_certificate_page(doc)
    doc = add_declaration_page(doc)
    doc = add_acknowledgment_page(doc)
    
    # Save the document
    output_file = "SMART_MARINE_REPORT_PART2.docx"
    doc.save(output_file)
    print(f"✅ Part 2 saved as {output_file}")

if __name__ == "__main__":
    main()
