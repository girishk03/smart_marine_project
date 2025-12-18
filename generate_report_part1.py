from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def create_document():
    doc = Document()
    
    # Set document margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.5)
    
    # Set default font and line spacing
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    return doc

def add_cover_page(doc):
    print("✅ Adding cover page...")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\nA Mini Project Report on\n\n')
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Add project title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SMART MARINE PROJECT:\n')
    run.font.size = Pt(20)
    run.font.bold = True
    run = p.add_run('AI-POWERED PLASTIC WASTE DETECTION AND\n')
    run.font.size = Pt(20)
    run.font.bold = True
    run = p.add_run('AUTONOMOUS COLLECTION SYSTEM\n\n\n')
    run.font.size = Pt(20)
    run.font.bold = True
    
    # Add submission details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted in partial fulfillment of the requirements\n')
    run.font.size = Pt(12)
    run = p.add_run('for the award of the degree of\n\n')
    run.font.size = Pt(12)
    
    # Add degree details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Bachelor of Technology\n\n')
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Add student details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('by\n\n')
    run.font.size = Pt(12)
    
    # Add student names
    students = [
        "SUDEEKSHA (20000XXXXX)",
        "[Student Name 2] (20000XXXXX)",
        "[Student Name 3] (20000XXXXX)",
        "[Student Name 4] (20000XXXXX)"
    ]
    
    for student in students:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{student}\n')
        run.font.size = Pt(12)
    
    # Add guide details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nUnder the guidance of\n\n')
    run.font.size = Pt(12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('[Guide Name]\n[Designation]\nDepartment of CSE\n\n\n')
    run.font.size = Pt(12)
    run.font.bold = True
    
    # Add department and university details
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING\n')
    run.font.size = Pt(14)
    run.font.bold = True
    run = p.add_run('ANURAG UNIVERSITY\n')
    run.font.size = Pt(14)
    run.font.bold = True
    run = p.add_run('VENKATAPUR-500088, TELANGANA\n')
    run.font.size = Pt(14)
    run.font.bold = True
    run = p.add_run(f'MONTH, {datetime.datetime.now().year}')
    run.font.size = Pt(12)
    
    # Add page break after cover page
    doc.add_page_break()
    
    return doc

def main():
    print("📄 Starting report generation (Part 1)...")
    doc = create_document()
    doc = add_cover_page(doc)
    
    # Save the document
    output_file = "SMART_MARINE_REPORT_PART1.docx"
    doc.save(output_file)
    print(f"✅ Part 1 saved as {output_file}")

if __name__ == "__main__":
    main()
